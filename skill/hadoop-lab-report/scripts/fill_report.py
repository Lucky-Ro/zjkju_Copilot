#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把内容填进学校实验报告模板(确定性 docx 组装器)。

引擎:python-docx 直接编辑**现有** template.docx(开→改单元格→存),保留官方版式;
代码块用淡金 #FFF2CC + ShadingType.CLEAR + 细黑边框 + 宽度 auto + Consolas 五号(长行 9pt);截图左对齐满栏、无图注。

两种用法:
  # 1) 由 Claude 产出 report.json(各栏目内容块),最忠实:
  python fill_report.py --config lab_config.json --template <skill>/assets/template.docx \
      --content runs/e04/report.json -o runs/e04/report.docx
  # 2) --auto:从 plan.json/run.log/state.json/shots 自动生成默认内容(兜底,可再人工微调)
  python fill_report.py --config lab_config.json --template <skill>/assets/template.docx \
      --auto --plan runs/e04/plan.json --shots runs/e04/shots --state runs/e04/state.json \
      -o runs/e04/report.docx

report.json 结构见本目录 README / references/report-template.md。
"""
from __future__ import annotations
import argparse
import json
import os
import re
import sys

import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from _common import load_config, apply_sid, eprint  # noqa: E402

# 样例排版指纹(实测自 风格参考.docx,固化进生成器——脱离样例文件也能复现,见 references/report-template.md)
CODE_FILL = "FFF2CC"        # 淡金
CODE_BORDER = "000000"      # 代码块边框=细黑 sz4 #000000(匹配样例,非旧的淡金 D9C27A)
CODE_FONT = "Consolas"
CODE_SIZE = Pt(10.5)        # 五号
CODE_SIZE_SMALL = Pt(9)     # 超长行降 9pt,免撑破(匹配样例对 sftp/put 长行的处理)
CODE_LONG = 96             # 行长超过该字符数视为「超长行」,降到 9pt
TABLE_W = 8300              # DXA,仅个别兜底用;代码块已改 auto 宽,不再据此定宽
IMG_W = Cm(15.5)            # 截图满栏(匹配样例 ≈15.49cm),左对齐等比
BODY_SIZE = 12              # 小四,正文
LINE = 1.0                  # 单倍行距(匹配样例;去掉 1.5 的「AI 味」留白)

LABELS = {
    "purpose": "实验目的",
    "materials": "实验内容及实验器材",
    "process": "实验过程",
    "conclusion": "实验结论",
    "summary": "实验总结",
}


# ───────── oxml 小工具 ─────────
def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    el.append(e)
    return e


def shade(cell, fill):
    _set(cell._tc.get_or_add_tcPr(), "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def set_table_auto(tbl):
    """代码块表宽 = auto(随内容):短命令小框、长命令撑满,匹配样例。
    细黑边框 sz4 #000000。不设 tblLayout fixed、不强制列宽(auto 模式由内容决定)。"""
    tblPr = tbl._tbl.tblPr
    _set(tblPr, "w:tblW", **{"w:w": "0", "w:type": "auto"})
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(borders, f"w:{side}", **{"w:val": "single", "w:sz": "4", "w:color": CODE_BORDER})
    tblPr.append(borders)


# ───────── 单元格内容操作 ─────────
def clear_keep_label(cell):
    """保留第一段(栏目标签),删除其余模板提示段落。"""
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def trim_trailing_empty(cell):
    """续写前裁掉单元格末尾的空段落(用户在占位标题下留的大段空白),
    使追加内容紧跟在已有内容之后,不留突兀的大空白。保留含文字/图片的段落。"""
    for p in reversed(cell.paragraphs[1:]):
        has_img = p._element.findall('.//' + qn('w:drawing'))
        if p.text.strip() == "" and not has_img:
            p._element.getparent().remove(p._element)
        else:
            break


def add_para(cell, text, size=BODY_SIZE, bold=False, italic=False, color=None,
             indent=False, space_before=0, space_after=6, line=LINE):
    # 默认左对齐(匹配样例,不居中);说明句无首行缩进、单倍行距
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Pt(size * 2)   # 首行缩进 2 个字(中文文档习惯;说明句不用)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    # 显式宋体(eastAsia),防个别环境中文回退;Latin 维持 docDefault(Times New Roman),与样例一致
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def _tiny_para(cell):
    """OOXML 要求嵌套表后跟一个段落;把它压到几乎不可见,避免代码块下方多出空行。"""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)
    return p


def add_code(cell, code, lang=None, tight_below=False):
    """代码块 = 单格嵌套表:#FFF2CC CLEAR 底 + Consolas 五号 + 细黑边框 + 宽度 auto(随内容)。
    超长行降 9pt 免撑破。tight_below=True 时(后面紧跟该命令的截图)下方不留空隙,做到「代码—截图」无缝。"""
    pre = cell.add_paragraph()
    pre.paragraph_format.space_before = Pt(2)
    pre.paragraph_format.space_after = Pt(0)
    pre.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pre.paragraph_format.line_spacing = Pt(3)
    nt = cell.add_table(rows=1, cols=1)
    set_table_auto(nt)
    nc = nt.rows[0].cells[0]
    shade(nc, CODE_FILL)
    _set(nc._tc.get_or_add_tcPr(), "w:tcW", **{"w:w": "0", "w:type": "auto"})  # 单元格也 auto,与表一致
    first = True
    for ln in (code.rstrip("\n").split("\n") or [""]):
        p = nc.paragraphs[0] if first else nc.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0       # 代码内单倍行距,紧凑但清晰
        r = p.add_run(ln)
        r.font.name = CODE_FONT
        r.font.size = CODE_SIZE_SMALL if len(ln) > CODE_LONG else CODE_SIZE
        rFonts = r._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), CODE_FONT)   # 中文也用 Consolas,防回退变形
    post = _tiny_para(cell)                       # 必需空段,压到最小
    if not tight_below:
        post.paragraph_format.line_spacing = Pt(8)  # 非紧贴时给一点点下方间距


def add_image(cell, path, caption=None, tight_above=False):
    """截图:**左对齐**(默认,匹配样例)、满栏 ≈15.5cm、等比、**图下无图注**。
    caption 形参保留以兼容旧调用,但不再渲染任何图注(样例全程无图注;说明句在图上方交代)。"""
    if not path or not os.path.exists(path):
        add_para(cell, f"[缺图: {path}]", size=10, italic=True, color="C00000")
        return
    p = cell.add_paragraph()                                       # 默认左对齐,不设 alignment
    p.paragraph_format.space_before = Pt(0 if tight_above else 4)  # 紧跟代码时不留缝
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=IMG_W)


# ───────── 渲染一个栏目 ─────────
def render_blocks(cell, blocks, sid3):
    # 排版铁律(匹配样例):说明句(上)→ 代码块 → 截图(下),全程无图注;
    # 说明句无首行缩进、单倍行距、不加粗;只有子任务/小节标题加粗。
    for b in blocks:
        t = b.get("type")
        if t == "para":
            txt = apply_sid(b.get("text", ""), sid3)
            # 「任务4.2 …」这类子任务标题当作小标题:加粗、上方留白;其余正文不加粗
            is_head = bool(re.match(r"^\s*任务\s*[\d.]+", txt)) or b.get("heading")
            add_para(cell, txt, bold=is_head, indent=False,
                     space_before=10 if is_head else 0, space_after=6, line=LINE)
        elif t == "code":
            add_code(cell, apply_sid(b.get("code", ""), sid3), b.get("lang"))
        elif t == "image":
            add_image(cell, b.get("path"))
        elif t == "caption":
            # 样例无图注:降级为普通说明段(不再渲染居中灰斜体小字)
            add_para(cell, apply_sid(b.get("text", ""), sid3), indent=False,
                     space_after=6, line=LINE)
        elif t == "step":
            # 说明句:普通短句,无编号、无缩进、单倍、不加粗(只有 para 的小节标题加粗)
            head = apply_sid(b.get("text", ""), sid3)
            add_para(cell, head, bold=False, indent=False, space_before=8, space_after=3, line=LINE)
            # 顺序:说明句(上)→ 命令(代码块)→ 截图(下),无图注。代码与其截图之间无缝。
            has_code, has_img = bool(b.get("code")), bool(b.get("image"))
            if has_code:
                add_code(cell, apply_sid(b["code"], sid3), b.get("lang"), tight_below=has_img)
            if has_img:
                add_image(cell, b["image"], tight_above=has_code)


# ───────── 表头 ─────────
def _nows(s):
    return s.replace(" ", "").replace("　", "").replace("\t", "")


def set_value_after(table, label, value):
    lab = _nows(label)
    for row in table.rows:
        cells = row.cells
        for i, c in enumerate(cells):
            if lab in _nows(c.text):
                # 找该行后面第一个与标签格不同的格作为值格
                for j in range(i + 1, len(cells)):
                    if cells[j]._tc is not c._tc:
                        cells[j].text = value
                        for p in cells[j].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(12)
                        return True
    eprint(f"[!] 表头未找到标签: {label}")
    return False


def fill_header(doc, cfg, title):
    ident = cfg["identity"]
    t0 = doc.tables[0]
    set_value_after(t0, "课程名称", ident.get("course_name", "Hadoop集群部署与开发"))
    set_value_after(t0, "实验名称", title)
    set_value_after(t0, "实验地点", ident.get("location", ""))
    set_value_after(t0, "实验时间", ident.get("exam_time", ""))
    set_value_after(t0, "学生姓名", ident.get("name", ""))
    set_value_after(t0, "学号", ident.get("student_id", ""))
    set_value_after(t0, "学院", ident.get("college", ""))
    set_value_after(t0, "专业班级", ident.get("major_class", ""))
    set_value_after(t0, "指导教师", ident.get("instructor", ""))


def content_cell(doc, key):
    label = LABELS[key]
    for row in doc.tables[1].rows:
        for c in row.cells:
            if c.paragraphs and label in c.paragraphs[0].text:
                return c
    return None


# ───────── --auto 默认内容 ─────────
def auto_content(plan, shots_dir, state):
    issues = (state or {}).get("issues", [])
    sec = {"purpose": [], "materials": [], "process": [], "conclusion": [], "summary": []}
    for s in plan["subtasks"]:
        if s.get("purpose"):
            sec["purpose"].append({"type": "para", "text": f"任务{s['subtask_id']} {s['title']}:{s['purpose']}"})
        mat = " ".join(x for x in (s.get("environment"), s.get("resources")) if x)
        if mat:
            sec["materials"].append({"type": "para", "text": mat})
    for s in plan["subtasks"]:
        sec["process"].append({"type": "para", "text": f"任务{s['subtask_id']} {s['title']}"})
        for st in s["steps"]:
            if st["kind"] not in ("auto", "author"):
                continue
            img = os.path.join(shots_dir, f"step-{s['subtask_id']}_{st['idx']}.png")
            blk = {"type": "step", "text": st["text"]}   # 说明句承担「这张图证明了什么」,放图上方;无图注
            if os.path.exists(img):
                blk["image"] = img
            if st["code"]:
                blk["code"] = st["code"]
                blk["lang"] = st["lang"]
            sec["process"].append(blk)
    # 结论:取末子任务最后一张图 + 提示
    sec["conclusion"].append({"type": "para", "text": "实验运行结果见上,关键查询与输出如下。"})
    # 总结:真实问题 + 心得
    if issues:
        sec["summary"].append({"type": "para", "text": "实验遇到的问题及解决方法:"})
        for it in issues:
            sec["summary"].append({"type": "para", "text": f"问题:{it.get('symptom','')}  解决:{it.get('fix','')}"})
    else:
        sec["summary"].append({"type": "para", "text": "实验过程顺利,未遇到阻塞性问题。"})
    sec["summary"].append({"type": "para", "text": "心得体会:通过本次实验,加深了对相关组件部署与使用流程的理解。"})
    return {"sections": sec}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="lab_config.json")
    ap.add_argument("--template", help="空白模板,从零生成")
    ap.add_argument("--into", help="在【现有报告】基础上续写:保留已填内容(如4.1),实验过程追加、结论/总结替换,表头不动")
    ap.add_argument("--content")
    ap.add_argument("--auto", action="store_true")
    ap.add_argument("--plan")
    ap.add_argument("--shots", default="")
    ap.add_argument("--state")
    ap.add_argument("-o", "--out", required=True)
    args = ap.parse_args()

    cfg = load_config(args.config)
    sid3 = cfg["identity"].get("student_id_last3", "")
    plan = json.load(open(args.plan, encoding="utf-8")) if args.plan else None
    title = (plan or {}).get("title") or "实验报告"

    if args.content:
        content = json.load(open(args.content, encoding="utf-8"))
    elif args.auto:
        state = json.load(open(args.state, encoding="utf-8")) if args.state and os.path.exists(args.state) else {}
        content = auto_content(plan, args.shots, state)
    else:
        eprint("需要 --content report.json 或 --auto(配合 --plan)")
        sys.exit(2)

    base = args.into or args.template
    if not base:
        eprint("需要 --template(从零)或 --into(在现有报告上续写)其一")
        sys.exit(2)
    into_mode = bool(args.into)
    append_sections = {"process"}  # 续写模式下「实验过程」保留已有(如4.1)、在末尾追加

    doc = docx.Document(base)
    if not into_mode:
        fill_header(doc, cfg, title)   # 续写模式表头已填,不动
    for key in ("purpose", "materials", "process", "conclusion", "summary"):
        blocks = content["sections"].get(key, [])
        if into_mode and not blocks:
            continue                    # 续写模式下,没给内容的栏目原样保留
        cell = content_cell(doc, key)
        if cell is None:
            eprint(f"[!] 模板里找不到栏目: {LABELS[key]}")
            continue
        if not (into_mode and key in append_sections):
            clear_keep_label(cell)      # 续写模式的「实验过程」不清,直接在4.1后追加
        else:
            trim_trailing_empty(cell)   # 追加前裁掉末尾空白,内容紧跟占位标题
        render_blocks(cell, blocks, sid3)

    os.makedirs(os.path.dirname(os.path.abspath(args.out)), exist_ok=True)
    doc.save(args.out)
    eprint(f"[OK] 已生成 {args.out}")


if __name__ == "__main__":
    main()
